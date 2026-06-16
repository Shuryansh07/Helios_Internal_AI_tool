"""Shared helpers for the meeting document (.docx) generators."""
import os
import re
from datetime import date


ACCENT = (0x1F, 0x3A, 0x5F)   # navy, matches the proposal docx builder


def safe_filename(title: str) -> str:
    base = re.sub(r"[^\w\- ]", "", title or "document").strip().replace(" ", "_")
    return (base or "document")[:60]


def out_path(doc_type: str, title: str, output_dir: str) -> str:
    os.makedirs(output_dir, exist_ok=True)
    fname = f"{doc_type}_{safe_filename(title)}_{date.today().isoformat()}.docx"
    return os.path.join(output_dir, fname)


def new_document(company: dict, doc_type_label: str, title: str, meta_bits=None):
    """Create a branded Document with a cover header. Returns (doc, helpers)."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    accent = RGBColor(*ACCENT)
    doc = Document()

    # Company line
    top = doc.add_paragraph()
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = top.add_run(company.get("name", "Company"))
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = accent

    # Document type + title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"{doc_type_label}\n{title}")
    r.bold = True
    r.font.size = Pt(15)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bits = [date.today().strftime("%d %B %Y")] + list(meta_bits or [])
    meta.add_run("  |  ".join(b for b in bits if b))
    doc.add_paragraph()

    def h1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = accent
        return p

    def h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        return p

    def para(text):
        return doc.add_paragraph(str(text)) if text else None

    def bullets(items, style="List Bullet"):
        for it in items or []:
            doc.add_paragraph(str(it), style=style)

    def numbered(items):
        for it in items or []:
            doc.add_paragraph(str(it), style="List Number")

    helpers = {"h1": h1, "h2": h2, "para": para, "bullets": bullets, "numbered": numbered}
    return doc, helpers


def footer_contact(doc, company: dict):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    contact = " | ".join(
        x for x in [company.get("email"), company.get("phone"), company.get("website")] if x
    )
    if contact:
        doc.add_paragraph()
        f = doc.add_paragraph(contact)
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
