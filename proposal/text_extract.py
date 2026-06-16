"""Extract plain text from uploaded files (.docx, .pdf, .txt, .xlsx, .csv) for
use as client requirement or reference-project input."""
import io
import os

SUPPORTED = {".docx", ".pdf", ".txt", ".xlsx", ".csv"}


def extract_text(filename: str, data: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".txt":
        return data.decode("utf-8", errors="ignore").strip()
    if ext == ".csv":
        return data.decode("utf-8", errors="ignore").strip()
    if ext == ".docx":
        return _from_docx(data)
    if ext == ".pdf":
        return _from_pdf(data)
    if ext == ".xlsx":
        return _from_xlsx(data)
    if ext == ".xls":
        raise ValueError(
            "Old .xls is not supported. Open it in Excel and Save As .xlsx, then re-upload."
        )
    raise ValueError(f"Unsupported file type '{ext}'. Use .docx, .pdf, .txt, .xlsx or .csv.")


def _from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _from_xlsx(data: bytes) -> str:
    """
    Render each sheet as a pipe-separated table with a sheet header.
    The format below is the most reliable way for LLMs to read tabular data —
    they parse it like a markdown table. Formulas are evaluated to their last
    cached value (data_only=True). Each sheet is capped to keep prompts sane.
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    MAX_ROWS_PER_SHEET = 500
    out = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            vals = ["" if v is None else str(v).strip() for v in row]
            while vals and not vals[-1]:
                vals.pop()
            if any(v for v in vals):
                rows.append(" | ".join(vals))
            if len(rows) >= MAX_ROWS_PER_SHEET:
                rows.append("... (sheet truncated)")
                break
        if rows:
            out.append(f"=== Sheet: {sheet_name} ===")
            out.extend(rows)
            out.append("")
    return "\n".join(out).strip()
