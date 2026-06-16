"""
Builds the Word (.docx) file.

Two modes:
  A) If you drop a Word template at template_docx/template.docx that contains
     Jinja-style placeholders (e.g. {{ project_title }}, {% for s in scope %}...),
     we render YOUR template with the AI content -> exact match to your format.
  B) Otherwise we generate a clean, professional default layout from scratch.
"""
import os
import re
from datetime import date


def _safe_filename(title: str) -> str:
    base = re.sub(r"[^\w\- ]", "", title or "proposal").strip().replace(" ", "_")
    return (base or "proposal")[:60]


# Pretty provider names used as the filename prefix.
_PROVIDER_LABELS = {"openai": "OpenAI", "gemini": "Gemini", "claude": "Claude"}


def build_docx(content: dict, company: dict, output_dir: str,
               template_path: str | None = None, provider: str | None = None) -> str:
    os.makedirs(output_dir, exist_ok=True)
    prefix = _PROVIDER_LABELS.get((provider or "").lower(), "")
    prefix = f"{prefix}_" if prefix else ""
    filename = f"{prefix}Proposal_{_safe_filename(content.get('project_title'))}_{date.today().isoformat()}.docx"
    out_path = os.path.join(output_dir, filename)

    if template_path and os.path.exists(template_path):
        _render_with_template(content, company, template_path, out_path)
    else:
        _build_default(content, company, out_path)
    return out_path


# ---------------------------------------------------------------------------
# Mode A: render the user's own Word template (matches their format exactly)
# ---------------------------------------------------------------------------
def _render_with_template(content, company, template_path, out_path):
    import jinja2
    from docxtpl import DocxTemplate

    doc = DocxTemplate(template_path)
    context = dict(content)
    context["company"] = company
    context["date"] = date.today().strftime("%d %B %Y")
    # autoescape=True so characters like "&" survive (both in our values and in
    # the template's own static text such as "Effort & Cost").
    doc.render(context, jinja_env=jinja2.Environment(autoescape=True))
    doc.save(out_path)


# ---------------------------------------------------------------------------
# Mode B: clean default proposal layout
# ---------------------------------------------------------------------------
def _build_default(content, company, out_path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    accent = RGBColor(0x1F, 0x3A, 0x5F)

    def heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = accent
        return p

    def bullets(items):
        for item in items or []:
            if isinstance(item, dict):
                item = " - ".join(str(v) for v in item.values() if v)
            doc.add_paragraph(str(item), style="List Bullet")

    # Cover / title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(company.get("name", "Project Proposal"))
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = accent
    if company.get("tagline"):
        sub = doc.add_paragraph(company["tagline"])
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pt.add_run("\nProject Proposal\n" + (content.get("project_title") or ""))
    r.bold = True
    r.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bits = [date.today().strftime("%d %B %Y")]
    if content.get("client_name"):
        bits.append("Prepared for: " + content["client_name"])
    meta.add_run("  |  ".join(bits))
    doc.add_page_break()

    heading("Objective")
    doc.add_paragraph(content.get("objective", ""))

    for ms in content.get("milestones", []):
        if isinstance(ms, dict):
            heading(ms.get("title", "Milestone"))
            bullets(ms.get("scope"))
        else:
            heading(str(ms))

    heading("Effort & Cost")
    doc.add_paragraph("Total Estimated Effort: " + (content.get("total_effort") or "To be confirmed"))
    doc.add_paragraph("Total Project Cost: " + (content.get("total_cost") or "To be confirmed"))

    # Footer contact line
    doc.add_paragraph()
    contact = " | ".join(
        x for x in [company.get("email"), company.get("phone"), company.get("website")] if x
    )
    if contact:
        f = doc.add_paragraph(contact)
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
